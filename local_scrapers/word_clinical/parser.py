import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .models import Estudio, Proclama, Referencia, SustentoProclamasDoc

# Palabras clave reconocidas como tipo_soporte cuando aparecen solas en la primera línea
_TIPOS_SOPORTE = {
    "sensorial", "instrumental", "seguridad", "formulación", "formulacion",
    "nombre de la marca", "nombre de la variedad", "clínico", "clinico",
}

_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _tc_text(tc) -> str:
    paragraphs = tc.findall(".//" + qn("w:p"))
    return "\n".join(
        "".join(r.text or "" for r in p.findall(".//" + qn("w:t")))
        for p in paragraphs
    ).strip()


def _cell_texts(row) -> list[str]:
    """
    Extract text from each column in a table row.
    Handles both plain <w:tc> cells and <w:sdt>-wrapped cells
    (Word content controls: dropdowns, date pickers).
    """
    result = []
    for child in row._tr:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "tc":
            result.append(_tc_text(child))
        elif local == "sdt":
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                for tc in sdt_content.findall(qn("w:tc")):
                    result.append(_tc_text(tc))
    return result


def _split_tipo_soporte(raw: str) -> tuple[str | None, str | None]:
    """
    Split a combined 'tipo\\nsoporte' cell into (tipo_soporte, soporte).
    Rules (in order):
      1. First line matches a known keyword → tipo / rest
      2. No newline, no period, short (<= 60 chars) → tipo only (label, not evidence)
      3. Otherwise → soporte only
    """
    if not raw:
        return None, None
    lines = raw.split("\n", 1)
    first = lines[0].strip()

    if first.lower() in _TIPOS_SOPORTE:
        soporte = lines[1].strip() if len(lines) > 1 else None
        return first, soporte or None

    # Short label with no period and not starting like a sentence → tipo_soporte
    # (e.g. "Nombre de la marca y de la variedad")
    _sentence_starters = (
        "el ", "la ", "los ", "las ", "se ", "en ", "es ", "un ", "una ",
        "the ", "this ", "for ", "over ", "at ", "in ",
    )
    starts_like_sentence = raw.lower().startswith(_sentence_starters)
    if "\n" not in raw and "." not in raw and len(raw) <= 60 and not starts_like_sentence:
        return raw.strip(), None

    return None, raw or None


def _parse_header(table) -> tuple[str | None, str | None]:
    """Tabla 0: extract fecha and version."""
    fecha = version = None
    for row in table.rows:
        cells = _cell_texts(row)
        for i, cell in enumerate(cells):
            if cell.strip() == "Fecha:" and i + 1 < len(cells):
                fecha = cells[i + 1].strip() or None
            if cell.strip() == "Versión:" and i + 1 < len(cells):
                version = cells[i + 1].strip() or None
    return fecha, version


def _parse_product_info(table) -> dict:
    """Tabla 1: key-value rows, skip merged header row."""
    info: dict[str, str] = {}
    for row in table.rows:
        cells = _cell_texts(row)
        if len(cells) < 2:
            continue
        key = cells[0].rstrip(":").strip()
        value = cells[1].strip()
        if key and value and key != value:  # skip merged header row
            info[key] = value
    return info


def _parse_estudios(table) -> list[Estudio]:
    estudios = []
    headers = None
    for row in table.rows:
        cells = _cell_texts(row)
        if not any(cells):
            continue
        if headers is None:
            headers = [c.strip() for c in cells]
            continue
        def get(col: str) -> str | None:
            try:
                idx = headers.index(col)
                val = cells[idx].strip() if idx < len(cells) else ""
                return val or None
            except ValueError:
                return None
        codigo = get("CÓDIGO") or get("CODIGO")
        nombre = get("NOMBRE")
        if not codigo and not nombre:
            continue
        estudios.append(Estudio(
            codigo=codigo or "",
            nombre=nombre or "",
            tipo_estudio=get("TIPO ESTUDIO"),
            agencia=get("AGENCIA"),
            fecha=get("FECHA"),
            informe=get("INFORME"),
        ))
    return estudios


def _parse_proclamas(table) -> list[Proclama]:
    proclamas = []
    is_header = True
    for row in table.rows:
        cells = _cell_texts(row)
        if is_header:
            is_header = False
            continue
        if not cells:
            continue
        proclama_text = cells[0].strip() if cells else ""
        if not proclama_text:
            continue

        # col1 may be tipo only, col2 may be soporte — or col1 has both merged
        if len(cells) >= 3 and cells[2].strip():
            tipo_soporte = cells[1].strip() or None
            soporte = cells[2].strip() or None
        else:
            raw = cells[1].strip() if len(cells) > 1 else ""
            tipo_soporte, soporte = _split_tipo_soporte(raw)

        proclamas.append(Proclama(
            proclama=proclama_text,
            tipo_soporte=tipo_soporte,
            soporte=soporte,
        ))
    return proclamas


def _parse_referencias(table) -> list[Referencia]:
    referencias = []
    for row in table.rows:
        cells = _cell_texts(row)
        if len(cells) < 2:
            continue
        numero = cells[0].strip()
        descripcion = cells[1].strip()
        if not numero or not descripcion:
            continue
        # skip header row
        if numero == "1" or numero.isdigit():
            referencias.append(Referencia(numero=numero, descripcion=descripcion))
        elif descripcion and numero not in ("", "CÓDIGO", "N°"):
            referencias.append(Referencia(numero=numero, descripcion=descripcion))
    return referencias


def _parse_intro(doc: Document) -> str | None:
    """Collect paragraphs under INTRODUCCIÓN section."""
    collecting = False
    parts = []
    stops = {"ESTUDIOS SEGURIDAD Y EFICACIA", "PLAN DE SUSTENTO", "BIBLIOGRAFÍA"}
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "INTRODUCCIÓN":
            collecting = True
            continue
        if text in stops:
            break
        if collecting and text:
            parts.append(text)
    return "\n\n".join(parts) or None


def parse_document(source: io.BytesIO | str | Path, filename: str) -> SustentoProclamasDoc:
    doc = Document(source)

    if len(doc.tables) < 5:
        raise ValueError(f"{filename}: se esperaban 5 tablas, encontradas {len(doc.tables)}")

    fecha, version = _parse_header(doc.tables[0])
    product_info = _parse_product_info(doc.tables[1])

    return SustentoProclamasDoc(
        archivo=filename,
        codigo_formula=product_info.get("Código fórmula") or product_info.get("Codigo formula") or "",
        nombre_producto=product_info.get("Nombre") or "",
        fecha_documento=fecha,
        version=version,
        modo_uso=product_info.get("Modo de uso"),
        precauciones=product_info.get("Precauciones"),
        intro_texto=_parse_intro(doc),
        estudios=_parse_estudios(doc.tables[2]),
        proclamas=_parse_proclamas(doc.tables[3]),
        referencias=_parse_referencias(doc.tables[4]),
    )
